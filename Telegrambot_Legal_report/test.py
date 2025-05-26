from docx import Document
from docx.shared import Pt
import os
from extraction import extract_org_data

def process_template(template_path, output_path, org_data):
    """Прямая запись данных в конкретные ячейки таблицы"""
    print(f"\n{'=' * 50}\nНачало обработки шаблона\n{'=' * 50}")

    # Проверка файла шаблона
    if not os.path.exists(template_path):
        print(f"❌ ОШИБКА: Файл шаблона не найден: {template_path}")
        return False

    try:
        doc = Document(template_path)
        print("✔ Шаблон успешно загружен")
    except Exception as e:
        print(f"❌ ОШИБКА при открытии шаблона: {str(e)}")
        return False

    # Проверка данных
    print("\nПроверка данных:")
    for key, value in org_data.items():
        print(f"{key}: {'✔' if value else '❌'} {value or 'Не найдено'}")

    # Обработка таблицы
    if not doc.tables:
        print("❌ В документе нет таблиц!")
        return False

    table = doc.tables[0]  # Первая таблица
    print(f"\nОбработка таблицы: {len(table.rows)} строк, {len(table.columns)} столбцов")

    try:
        # Ячейка Наименование (строка 1, столбец 2)
        name_cell = table.cell(0, 1)  # Индексы с 0
        name_cell.text = org_data['Организация'] or 'Наименование не найдено'
        print(f"✔ Записано наименование в ячейку 1.2: {name_cell.text}")

        # Ячейка ОГРН (строка 2, столбец 2)
        ogrn_cell = table.cell(1, 1)  # Индексы с 0
        ogrn_cell.text = org_data['ОГРН'] or 'ОГРН не найден'
        print(f"✔ Записан ОГРН в ячейку 2.2: {ogrn_cell.text}")

        # Ячейка ИНН/КПП (строка 3, столбец 2)
        inn_kpp_cell = table.cell(2, 1)
        replacement = f"{org_data['ИНН'] or 'ИНН не найден'}/{org_data['КПП'] or 'КПП не найден'}"
        inn_kpp_cell.text = replacement
        print(f"✔ Записан ИНН/КПП в ячейку 3.2: {inn_kpp_cell.text}")

    except IndexError:
        print("❌ ОШИБКА: Неправильная структура таблицы!")
        return False

    # Сохранение
    try:
        doc.save(output_path)
        print(f"\n✔ Файл успешно сохранен: {output_path}")
        print(f"Размер файла: {os.path.getsize(output_path)} байт")
        return True
    except Exception as e:
        print(f"❌ ОШИБКА при сохранении: {str(e)}")
        return False

if __name__ == "__main__":
    source_file = 'контур_фокус.docx'
    template_file = 'шаблон.docx'
    output_file = 'НОВЫЙ_ФАЙЛ.docx'

    print("=" * 50)
    print("Начало обработки документов")
    print("=" * 50)

    # Извлекаем данные
    org_data = extract_org_data(source_file)

    # Обрабатываем шаблон
    success = process_template(template_file, output_file, org_data)