from docx import Document
from docx.shared import Parented
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re


def is_strikethrough(run):
    """Проверяет, является ли текст зачеркнутым."""
    if run.font.strike:
        return True
    rPr = run._element.get_or_add_rPr()
    strike = rPr.find(qn('w:strike'))
    if strike is not None and strike.get(qn('w:val')) != '0':
        return True
    return False


def clean_sum_text(sum_text):
    """Приводит сумму к формату '1000000,00' (без пробелов, символов валюты и прочего текста)."""
    if not sum_text:
        return ''
    cleaned = sum_text.replace('\xa0', '').replace(' ', '')
    cleaned = re.sub(r'руб(\.|ль)?|р\.|₽', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'[^\d,]', '', cleaned)
    cleaned = cleaned.replace('.', ',')
    if cleaned.count(',') > 1:
        first = cleaned.find(',')
        cleaned = cleaned[:first + 1] + cleaned[first + 1:].replace(',', '')
    return cleaned


def extract_text_without_strikethrough(paragraph):
    """Извлекает текст из параграфа, исключая зачеркнутые части."""
    return ''.join(run.text for run in paragraph.runs if not is_strikethrough(run))


def extract_table_text_without_strikethrough(table):
    """Извлекает текст из таблицы, исключая зачеркнутые части."""
    text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = []
            for paragraph in cell.paragraphs:
                cell_text.append(extract_text_without_strikethrough(paragraph))
            row_text.append(' '.join(cell_text))
        text.append(' | '.join(row_text))
    return '\n'.join(text)


def extract_basic_info(doc):
    """Извлекает основную информацию о компании."""
    basic_info = {
        'Краткое наименование': '',
        'ИНН': '',
        'КПП': '',
        'ОГРН': '',
        'Дата образования': '',
        'Юридический адрес': '',
        'Уставный капитал': '',
        'Генеральный директор': '',
        'ОКВЭД(основной)': ''
    }

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip() if row.cells[
                0].paragraphs else ''
            value = extract_text_without_strikethrough(row.cells[1].paragraphs[0]).strip() if row.cells[
                1].paragraphs else ''

            if key == "Краткое наименование":
                basic_info['Краткое наименование'] = value
            elif key == "ИНН":
                basic_info['ИНН'] = value
            elif key == "КПП":
                basic_info['КПП'] = value
            elif key == "ОГРН":
                basic_info['ОГРН'] = value
            elif key == "Дата образования":
                basic_info['Дата образования'] = value
            elif key in ["Юр. адрес", "Юридический адрес"]:
                basic_info['Юридический адрес'] = value
            elif key == "Уставный капитал":
                basic_info['Уставный капитал'] = value
            elif key == "Генеральный директор":
                basic_info['Генеральный директор'] = value
            elif key == "Основной вид деятельности":
                basic_info['ОКВЭД(основной)'] = value

    return basic_info


def extract_staff_info(doc):
    """Извлекает информацию о сотрудниках."""
    staff_info = {
        'Среднесписочная численность': {'year_1': '', 'year_2': '', 'year_3': ''},
        'Расходы на оплату труда': {'year_1': '', 'year_2': '', 'year_3': ''}
    }
    staff_years = {}

    # Извлечение данных о численности
    for table in doc.tables:
        rows = table.rows
        i = 0
        while i < len(rows):
            cells = rows[i].cells
            key = extract_text_without_strikethrough(cells[0].paragraphs[0]).strip() if cells[0].paragraphs else ''
            value = extract_text_without_strikethrough(cells[1].paragraphs[0]).strip() if len(cells) > 1 and cells[
                1].paragraphs else ''

            if 'Среднесписоч' in key:
                text_block = value
                j = i + 1
                while j < len(rows):
                    next_key = extract_text_without_strikethrough(rows[j].cells[0].paragraphs[0]).strip() if \
                    rows[j].cells[0].paragraphs else ''
                    next_value = extract_text_without_strikethrough(rows[j].cells[1].paragraphs[0]).strip() if len(
                        rows[j].cells) > 1 and rows[j].cells[1].paragraphs else ''

                    if next_key != '':
                        break
                    text_block += "\n" + next_value
                    j += 1

                matches = re.findall(r'за (\d{4}):\s*([0-9]+)', text_block)
                for year, val in matches:
                    staff_years[year] = val
                i = j
                continue
            i += 1

    # Извлечение данных о расходах на оплату труда
    for table in doc.tables:
        header_row = None
        for row in table.rows:
            row_values = [extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower()
                          for cell in row.cells if cell.paragraphs]
            if (any('код' in cell for cell in row_values) and
                    sum(re.search(r'20\d{2}', cell) is not None for cell in row_values) >= 2):
                header_row = row
                break

        if header_row:
            year_indices = {}
            for idx, cell in enumerate(header_row.cells):
                if not cell.paragraphs:
                    continue
                match = re.search(r'(20\d{2})', extract_text_without_strikethrough(cell.paragraphs[0]))
                if match:
                    year_indices[match.group(1)] = idx

            for row in table.rows:
                first_cell = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower() if \
                row.cells[0].paragraphs else ''
                code_cell = extract_text_without_strikethrough(row.cells[1].paragraphs[0]).strip() if len(
                    row.cells) > 1 and row.cells[1].paragraphs else ''

                if (('оплата труда работников' in first_cell or code_cell == '4122') and
                        len(row.cells) > 2):
                    year_to_val = {}
                    for year, idx in year_indices.items():
                        if idx < len(row.cells) and row.cells[idx].paragraphs:
                            val = extract_text_without_strikethrough(row.cells[idx].paragraphs[0]).strip().replace(' ',
                                                                                                                   '')
                            year_to_val[year] = val

                    staff_years_sorted = sorted(staff_years, reverse=True)
                    slots = ['year_1', 'year_2', 'year_3']
                    for idx, year in enumerate(staff_years_sorted[:3]):
                        if year in year_to_val:
                            staff_info['Расходы на оплату труда'][slots[idx]] = f"{{'{year}': '{year_to_val[year]}'}}"
                    break

    # Заполнение данных о численности
    staff_years_sorted = sorted(staff_years, reverse=True)
    slots = ['year_1', 'year_2', 'year_3']
    for idx, year in enumerate(staff_years_sorted[:3]):
        val = staff_years[year]
        staff_info['Среднесписочная численность'][slots[idx]] = f"{{'{year}': '{val}'}}"

    return staff_info


def extract_founders(doc):
    """Извлекает информацию об учредителях."""
    founders = []
    for table in doc.tables:
        header = [extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower()
                  for cell in table.rows[0].cells if cell.paragraphs]

        col_map = {'share': None, 'sum': None, 'name': None, 'date': None}
        for idx, text in enumerate(header):
            if 'доля в %' in text or 'доля (%)' in text:
                col_map['share'] = idx
            elif 'доля в руб' in text or 'вклад' in text or 'сумма' in text:
                col_map['sum'] = idx
            elif 'наименование' in text or 'участник' in text or 'фио' in text:
                col_map['name'] = idx
            elif 'дата' in text:
                col_map['date'] = idx

        if col_map['share'] is not None and col_map['sum'] is not None and col_map['name'] is not None:
            i = 1
            while i < len(table.rows):
                row = table.rows[i]
                cells = row.cells
                get_cell = lambda idx: extract_text_without_strikethrough(cells[idx].paragraphs[0]).strip() \
                    if idx is not None and idx < len(cells) and cells[idx].paragraphs else ''

                share = get_cell(col_map['share'])
                summ = get_cell(col_map['sum'])
                name = get_cell(col_map['name'])
                date = get_cell(col_map['date']) if col_map['date'] is not None else ''

                if not date and i + 1 < len(table.rows):
                    next_row = table.rows[i + 1]
                    if (len(next_row.cells) > 0 and
                            re.match(r'\d{2}\.\d{2}\.\d{4}',
                                     extract_text_without_strikethrough(next_row.cells[0].paragraphs[0]).strip())):
                        date = extract_text_without_strikethrough(next_row.cells[0].paragraphs[0]).strip()
                        i += 1

                if share or summ or name:
                    founders.append({
                        "share": share,
                        "sum": clean_sum_text(summ),
                        "name": name,
                        "date": date,
                    })
                i += 1
    return founders


def extract_collaterals(doc):
    """Извлекает информацию о залогах."""
    collaterals = []
    all_text = []

    for para in doc.paragraphs:
        text = extract_text_without_strikethrough(para)
        if text.strip():
            all_text.append(text)

    for table in doc.tables:
        all_text.append(extract_table_text_without_strikethrough(table))

    text = '\n'.join(all_text)

    collateral_pattern = re.compile(
        r'(?P<date>\d{2}\.\d{2}\.\d{4})\s+(?P<regnum>\S+)\s*'
        r'(?:\s*Залогодатель[^\n]*\n(?P<pledger>.*?)\n)?'
        r'(?:\s*Залогодержатель[^\n]*\n(?P<holder>.*?)\n)?'
        r'(?:\s*Договор[^\n]*\n(?P<contract>.*?)\n)?'
        r'(?:\s*Срок исполнения[^\n]*\n(?P<term>.*?)\n)?'
        r'(?:\s*Тип имущества[^\n]*\n(?P<asset_type>.*?)\n)?'
        r'(?:\s*Описание[^\n]*\n(?P<asset>.*?))'
        r'(?:\n{2,}|$)',
        re.DOTALL | re.MULTILINE
    )

    for m in collateral_pattern.finditer(text):
        entry = {
            'date': m.group('date').strip() if m.group('date') else '',
            'regnum': m.group('regnum').strip() if m.group('regnum') else '',
            'pledger': m.group('pledger').strip() if m.group('pledger') else '',
            'holder': m.group('holder').strip() if m.group('holder') else '',
            'contract': m.group('contract').strip() if m.group('contract') else '',
            'term': m.group('term').strip() if m.group('term') else '',
            'asset_type': m.group('asset_type').strip() if m.group('asset_type') else '',
            'asset': m.group('asset').strip() if m.group('asset') else '',
        }
        entry['asset'] = re.split(r'\n\d{2}\.\d{2}\.\d{4}\s+\S+', entry['asset'])[0].strip()
        if entry['date'] and (entry['holder'] or entry['asset']):
            collaterals.append(entry)
    return collaterals


def extract_leasing_info(doc):
    """Извлекает информацию о лизинге."""
    leasing_info_all = []
    for table in doc.tables:
        if len(table.columns) != 2 or len(table.rows) < 5:
            continue

        first_cell_text = extract_text_without_strikethrough(table.cell(0, 0).paragraphs[0]).strip()
        if not re.match(r'\d{2}\.\d{2}\.\d{4}', first_cell_text):
            continue

        data = {
            "Лизингодатель": "",
            "Период лизинга": "",
            "Категория": "",
            "Текущий статус": ""
        }

        for row in table.rows:
            key = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower()
            value = extract_text_without_strikethrough(row.cells[1].paragraphs[0]).strip()

            if key.startswith("лизингодатель"):
                if "Сведения скрыты" in value:
                    data["Лизингодатель"] = "Сведения скрыты"
                else:
                    data["Лизингодатель"] = value.split("— ИНН")[0].strip() or value.strip()
            elif key.startswith("период лизинга"):
                data["Период лизинга"] = value
            elif key.startswith("категория"):
                data["Категория"] = value
            elif key.startswith("статус"):
                data["Текущий статус"] = value.split('\n')[0].strip()

        if data["Лизингодатель"] or data["Категория"]:
            leasing_info_all.append(data)

    not_finished = [
        x for x in leasing_info_all
        if 'завершился' not in x.get('Текущий статус', '').lower()
    ]
    if not leasing_info_all:
        return ["Информации по лизингу нет"]
    if not not_finished:
        return ["Информации по лизингу со статусом 'Действует' нет"]
    return not_finished


def extract_credit_debt(doc):
    """Извлекает информацию о кредиторской задолженности."""
    for table in doc.tables:
        headers = [extract_text_without_strikethrough(cell.paragraphs[0]).lower()
                   for cell in table.rows[0].cells if cell.paragraphs]
        if (len(headers) >= 4
                and any('код' in h for h in headers)
                and sum(re.search(r'20\d{2}', h) is not None for h in headers) >= 2):

            debt_row = None
            for row in table.rows:
                if extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower().startswith(
                        "кредиторская задолженность"):
                    debt_row = row
                    break

            if debt_row:
                year_val = {}
                for idx, cell in enumerate(headers):
                    year_match = re.search(r'(20\d{2})', cell)
                    if year_match and idx < len(debt_row.cells):
                        year = year_match.group(1)
                        value = extract_text_without_strikethrough(debt_row.cells[idx].paragraphs[0]).strip().replace(
                            ' ', '').replace('–', '0')
                        year_val[year] = value

                years_sorted = sorted(year_val.keys(), reverse=True)
                slots = ['year_1', 'year_2', 'year_3']
                res = {}
                for i, year in enumerate(years_sorted[:3]):
                    val = year_val[year]
                    res[slots[i]] = f"{{'{year}': '{val}'}}"
                return res
    return {'year_1': '', 'year_2': '', 'year_3': ''}


def extract_financial_results(doc):
    """Извлекает финансовые результаты."""
    financial_results = {}
    year_indices = {}
    last_years = []

    for table in doc.tables:
        header_row = table.rows[0]
        header_cells = [extract_text_without_strikethrough(cell.paragraphs[0]).strip()
                        for cell in header_row.cells if cell.paragraphs]
        temp_year_indices = {}

        for idx, text in enumerate(header_cells):
            match = re.match(r'конец\s*(20\d{2})', text.lower()) or re.match(r'(20\d{2})', text)
            if match:
                year = match.group(1)
                temp_year_indices[year] = idx

        if len(temp_year_indices) >= 2:
            for row in table.rows:
                first_cell = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower() if \
                row.cells[0].paragraphs else ''
                if 'выручка' in first_cell:
                    year_indices = temp_year_indices
                    break

        if year_indices:
            last_years = sorted(year_indices, reverse=True)[:4]
            for row in table.rows[1:]:
                if len(row.cells) < max(year_indices.values()) + 1:
                    continue

                row_name = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip() if row.cells[
                    0].paragraphs else ''
                if not row_name:
                    continue

                values = {}
                for y in last_years:
                    idx = year_indices[y]
                    if idx < len(row.cells) and row.cells[idx].paragraphs:
                        val = extract_text_without_strikethrough(row.cells[idx].paragraphs[0]).strip().replace(' ',
                                                                                                               '').replace(
                            '–', '0')
                        values[y] = val

                if any(values.values()):
                    financial_results[row_name] = values
            break

    return financial_results


def extract_assets_and_receivables(doc):
    """Извлекает информацию об основных средствах и дебиторской задолженности."""
    result = {
        "Основные средства": {},
        "Дебиторская задолженность": {}
    }

    for table in doc.tables:
        header_cells = [extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower()
                        for cell in table.rows[0].cells if cell.paragraphs]
        year_to_col = {}

        for idx, txt in enumerate(header_cells):
            m = re.search(r'конец\s*(20\d{2})', txt)
            if m:
                year_to_col[m.group(1)] = idx

        if len(year_to_col) >= 3:
            years_sorted = sorted(year_to_col.keys(), reverse=True)[:3]
            row_ос = None
            row_дебит = None

            for row in table.rows:
                first_cell = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower() if \
                row.cells[0].paragraphs else ''
                if first_cell == "основные средства":
                    row_ос = row
                if first_cell == "дебиторская задолженность":
                    row_дебит = row

            if row_ос and row_дебит:
                for i, year in enumerate(years_sorted):
                    idx = year_to_col[year]
                    val_os = extract_text_without_strikethrough(row_ос.cells[idx].paragraphs[0]).strip().replace(' ',
                                                                                                                 '').replace(
                        '–', '0') if idx < len(row_ос.cells) and row_ос.cells[idx].paragraphs else '0'
                    val_db = extract_text_without_strikethrough(row_дебит.cells[idx].paragraphs[0]).strip().replace(' ',
                                                                                                                    '').replace(
                        '–', '0') if idx < len(row_дебит.cells) and row_дебит.cells[idx].paragraphs else '0'
                    result["Основные средства"][f'year_{i + 1}'] = {year: val_os}
                    result["Дебиторская задолженность"][f'year_{i + 1}'] = {year: val_db}
                break

    return result


def parsing_all_docx(docx_path):
    """Основная функция для парсинга всего документа."""
    company_data = {
        'Краткое наименование': '',
        'ИНН': '',
        'КПП': '',
        'ОГРН': '',
        'Дата образования': '',
        'Юридический адрес': '',
        'Уставный капитал': '',
        'Генеральный директор': '',
        'Учредители/участники': [],
        'ОКВЭД(основной)': '',
        'Сведения о сотрудниках': {
            'Среднесписочная численность': {'year_1': '', 'year_2': '', 'year_3': ''},
            'Расходы на оплату труда': {'year_1': '', 'year_2': '', 'year_3': ''},
        },
        'Сведения о залогах': [],
        'Сведения о лизинге': [],
        'Кредиторская задолженность': {'year_1': '', 'year_2': '', 'year_3': ''},
        'Отчет о финансовых результатах': {},
        'Основные средства и дебиторка': {
            'Основные средства': {},
            'Дебиторская задолженность': {}
        }
    }

    if not os.path.isfile(docx_path):
        print(f"Файл '{docx_path}' не найден.")
        return company_data

    try:
        doc = Document(docx_path)

        # Извлечение данных
        basic_info = extract_basic_info(doc)
        staff_info = extract_staff_info(doc)
        founders = extract_founders(doc)
        collaterals = extract_collaterals(doc)
        leasing_info = extract_leasing_info(doc)
        credit_debt = extract_credit_debt(doc)
        financial_results = extract_financial_results(doc)
        assets_receivables = extract_assets_and_receivables(doc)

        # Объединение данных
        company_data.update(basic_info)
        company_data['Сведения о сотрудниках'] = staff_info
        company_data['Учредители/участники'] = founders
        company_data['Сведения о залогах'] = collaterals
        company_data['Сведения о лизинге'] = leasing_info
        company_data['Кредиторская задолженность'] = credit_debt
        company_data['Отчет о финансовых результатах'] = financial_results
        company_data['Основные средства и дебиторка'] = assets_receivables

    except Exception as e:
        print(f"Ошибка при обработке файла '{docx_path}': {e}")

    return company_data


if __name__ == "__main__":
    docx_path = os.path.join(os.path.dirname(__file__), "word.docx")
    result = parsing_all_docx(docx_path)
    from pprint import pprint

    pprint(result)