from docx import Document
from docx.shared import Parented
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pprint import pprint
import zipfile
import re
import xml.etree.ElementTree as ET
import docx2txt
from docx.table import Table
import os
from docx.table import _Cell


def extract_inn_ogrn(text):
    """Извлекает ИНН и ОГРН из строки и возвращает очищенный текст без ИНН/ОГРН, а также отдельные значения."""
    inn_match = re.search(r'\bИНН[\s:\xa0]*([0-9]{10,12})\b', text)
    ogrn_match = re.search(r'\bОГРН[\s:\xa0]*([0-9]{13})\b', text)
    inn = inn_match.group(1) if inn_match else ''
    ogrn = ogrn_match.group(1) if ogrn_match else ''
    cleaned_text = re.sub(r'ИНН[\s:\xa0]*[0-9]{10,12}', '', text)
    cleaned_text = re.sub(r'ОГРН[\s:\xa0]*[0-9]{13}', '', cleaned_text)
    return cleaned_text.strip(" ,–—\u2002"), inn, ogrn


def extract_text_from_cell(cell):
    """Извлекает и объединяет текст из всех непустых параграфов ячейки."""
    if not hasattr(cell, 'paragraphs'):
        return ''
    return '\n'.join(
        paragraph.text.strip()
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    ).strip()


def extract_text_without_strikethrough(paragraph):
    """Извлекает текст из параграфа, исключая зачеркнутые части."""
    if not hasattr(paragraph, 'runs'):
        return ''
    return ''.join(run.text for run in paragraph.runs if not is_strikethrough(run))


def is_strikethrough(run):
    """Проверяет, является ли текст зачеркнутым."""
    if hasattr(run.font, 'strike') and run.font.strike:
        return True
    rPr = run._element.get_or_add_rPr()
    strike = rPr.find(qn('w:strike'))
    return strike is not None and strike.get(qn('w:val')) != '0'


def extract_competitive_manager(doc):
    """Извлекает информацию о конкурсном управляющем из таблиц документа."""
    target_keys = [
        "Исполняющий обязанности конкурсного управляющего",
        "Конкурсный управляющий"
    ]

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = extract_text_from_cell(row.cells[0])
            value = extract_text_from_cell(row.cells[1])

            if any(target.lower() in key.lower() for target in target_keys):
                return split_director_info(value)

    return {'ФИО': '', 'ИНН': ''}


def split_director_info(text):
    """Возвращает словарь с ФИО и ИНН, очищает пробелы и неразрывные пробелы."""
    cleaned = re.sub(r'[\u00A0\s]', '', text)  # удаляем пробелы и неразрывные пробелы
    match = re.search(r'(\d{10}|\d{12})', cleaned)

    if match:
        inn = match.group(1)
        raw_fio = text[:text.find('ИНН')].strip(' ,')
        return {'ФИО': raw_fio, 'ИНН': inn}
    else:
        return {'ФИО': text.strip(), 'ИНН': ''}


def clean_sum_text(sum_text):
    """Приводит сумму к формату '1000000,00' (без пробелов, символов валюты и прочего текста)."""
    if not sum_text:
        return ''
    cleaned = sum_text.replace('\xa0', '').replace(' ', '')
    cleaned = re.sub(r'руб(\.|ль)?|р\.|₽', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'[^\d,]', '', cleaned)
    cleaned = cleaned.replace('.', ',')
    if cleaned.count(',') > 1:
        first = cleaned.find(',')
        cleaned = cleaned[:first + 1] + cleaned[first + 1:].replace(',', '')
    return cleaned


def extract_first_address_block(full_address: str) -> str:
    """
    Извлекает только первый логически завершённый адресный блок из полной строки.
    Предполагается, что новые блоки начинаются с нового индекса (например, 115280).
    """
    match = list(re.finditer(r'\b1\d{5}\b', full_address))  # ищем все индексы

    if len(match) >= 2:
        return full_address[:match[1].start()].rstrip(', ')
    return full_address.strip(', ')


def format_founders(founders):
    """Форматирует информацию об учредителях в читаемый вид"""
    formatted = []
    for idx, founder in enumerate(founders, 1):
        share = founder.get('Доля в %', '')
        name = founder.get('Наимен. и реквизиты', '')
        formatted.append(f"{idx}. {name} — {share}")
    return "\n".join(formatted)


def extract_basic_info(doc):
    basic_info = {
        'Наименование': '',
        'ИНН': '',
        'КПП': '',
        'ОГРН': '',
        'Дата создания': '',
        'Юридический адрес': [],
        'Уставный капитал': '',
        'Генеральный директор': {'ФИО': '', 'ИНН': ''},
        'ОКВЭД (основной)': '',
        'Учредители/участники (текущие)': []
    }

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = extract_text_from_cell(row.cells[0])
            value = extract_text_from_cell(row.cells[1])

            if not key and not value:
                continue

            if "Краткое наименование" in key:
                basic_info['Наименование'] = value
            elif key == "ИНН":
                basic_info['ИНН'] = value
            elif key == "КПП":
                basic_info['КПП'] = value
            elif key == "ОГРН":
                basic_info['ОГРН'] = value
            elif key == "Дата образования":
                basic_info['Дата создания'] = value
            elif "Юр. адрес" in key or "Юридический адрес" in key:
                if value:
                    basic_info['Юридический адрес'].extend(value.split('\n'))
            elif "Уставный капитал" in key and not basic_info['Уставный капитал']:
                basic_info['Уставный капитал'] = value
            elif "Генеральный директор" in key:
                basic_info['Генеральный директор'] = split_director_info(value)
            elif "Основной вид деятельности" in key:
                basic_info['ОКВЭД (основной)'] = value

    raw_address = ', '.join(
        addr.strip() for addr in basic_info['Юридический адрес'] if addr.strip())
    basic_info['Юридический адрес'] = extract_first_address_block(raw_address)

    return basic_info


def extract_staff_info(doc):
    """Извлекает информацию о сотрудниках: численность и среднюю зарплату."""
    staff_info = {
        'Среднесписочная численность': {'year_1': '', 'year_2': '', 'year_3': ''},
        'Средняя заработная плата': {'year_1': '', 'year_2': '', 'year_3': ''}
    }
    staff_years_count = {}
    staff_years_salary = {}

    for table in doc.tables:
        rows = table.rows
        i = 0
        while i < len(rows):
            if len(rows[i].cells) < 2:
                i += 1
                continue

            key_cell = rows[i].cells[0]
            value_cell = rows[i].cells[1] if len(rows[i].cells) > 1 else None

            key = ''
            if key_cell.paragraphs:
                key = extract_text_without_strikethrough(key_cell.paragraphs[0]).strip().lower()

            value = ''
            if value_cell and value_cell.paragraphs:
                value = extract_text_without_strikethrough(value_cell.paragraphs[0]).strip()

            if 'среднесписоч' in key:
                text_block_count = value
                j = i + 1
                while j < len(rows):
                    if len(rows[j].cells) < 2:
                        j += 1
                        continue

                    next_key_cell = rows[j].cells[0]
                    next_value_cell = rows[j].cells[1] if len(rows[j].cells) > 1 else None

                    next_key = ''
                    if next_key_cell.paragraphs:
                        next_key = extract_text_without_strikethrough(next_key_cell.paragraphs[0]).strip().lower()

                    next_value = ''
                    if next_value_cell and next_value_cell.paragraphs:
                        next_value = extract_text_without_strikethrough(next_value_cell.paragraphs[0]).strip()

                    if next_key:
                        break
                    text_block_count += "\n" + next_value
                    j += 1

                matches = re.findall(r'за\s+(\d{4}):\s*([0-9]+)', text_block_count)
                for year, val in matches:
                    staff_years_count[year] = val
                i = j
                continue

            elif 'средняя заработная плата' in key or 'среднемесячная заработная плата' in key:
                text_block_salary = value
                j = i + 1
                while j < len(rows):
                    if len(rows[j].cells) < 2:
                        j += 1
                        continue

                    next_key_cell = rows[j].cells[0]
                    next_value_cell = rows[j].cells[1] if len(rows[j].cells) > 1 else None

                    next_key = ''
                    if next_key_cell.paragraphs:
                        next_key = extract_text_without_strikethrough(next_key_cell.paragraphs[0]).strip().lower()

                    next_value = ''
                    if next_value_cell and next_value_cell.paragraphs:
                        next_value = extract_text_without_strikethrough(next_value_cell.paragraphs[0]).strip()

                    if next_key:
                        break
                    text_block_salary += "\n" + next_value
                    j += 1

                matches = re.findall(r'за\s+(\d{4}):\s*([\d\s]+)', text_block_salary)
                for year, val in matches:
                    salary_clean = clean_sum_text(val)
                    staff_years_salary[year] = salary_clean
                i = j
                continue

            i += 1

    slots = ['year_1', 'year_2', 'year_3']
    for idx, year in enumerate(sorted(staff_years_count.keys(), reverse=True)[:3]):
        staff_info['Среднесписочная численность'][slots[idx]] = f"{{'{year}': '{staff_years_count[year]}'}}"

    for idx, year in enumerate(sorted(staff_years_salary.keys(), reverse=True)[:3]):
        staff_info['Средняя заработная плата'][slots[idx]] = f"{{'{year}': '{staff_years_salary[year]}'}}"

    return staff_info


def extract_founders(doc):
    """Извлекает информацию об учредителях"""
    actual = []
    outdated = []

    for table in doc.tables:
        if len(table.rows) == 0:
            continue

        header = []
        for cell in table.rows[0].cells:
            if cell.paragraphs and len(cell.paragraphs) > 0:
                header.append(extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower())
            else:
                header.append('')

        col_map = {'share': None, 'sum': None, 'name': None, 'date': None}
        for idx, text in enumerate(header):
            if 'доля' in text and '%' in text:
                col_map['share'] = idx
            elif 'руб' in text or 'вклад' in text or 'сумма' in text:
                col_map['sum'] = idx
            elif 'участник' in text or 'наименование' in text or 'фио' in text:
                col_map['name'] = idx
            elif 'дата' in text:
                col_map['date'] = idx

        if not all(col_map[k] is not None for k in ('share', 'sum', 'name')):
            continue

        i = 1
        while i < len(table.rows):
            row = table.rows[i]
            if len(row.cells) == 0:
                i += 1
                continue

            def get_full_text(cell):
                if not hasattr(cell, 'paragraphs'):
                    return ''
                return '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()

            def has_strikethrough(cell):
                if not hasattr(cell, 'paragraphs'):
                    return False
                for p in cell.paragraphs:
                    for run in p.runs:
                        if is_strikethrough(run) and run.text.strip():
                            return True
                return False

            share = get_full_text(row.cells[col_map['share']]) if col_map['share'] < len(row.cells) else ''
            summ = get_full_text(row.cells[col_map['sum']]) if col_map['sum'] < len(row.cells) else ''
            name_cell = row.cells[col_map['name']] if col_map['name'] < len(row.cells) else None
            name = get_full_text(name_cell) if name_cell else ''
            has_strike = has_strikethrough(name_cell) if name_cell else False
            date = get_full_text(row.cells[col_map['date']]) if col_map['date'] is not None and col_map['date'] < len(
                row.cells) else ''

            if not date and i + 1 < len(table.rows):
                next_row = table.rows[i + 1]
                if len(next_row.cells) > 0 and next_row.cells[0].paragraphs:
                    candidate = extract_text_without_strikethrough(next_row.cells[0].paragraphs[0]).strip()
                    if re.match(r'\d{2}\.\d{2}\.\d{4}', candidate):
                        date = candidate
                        i += 1

            name_clean, inn, _ = extract_inn_ogrn(name)
            full_name = f"{name_clean}, ИНН {inn}".strip(', ') if inn else name_clean

            record = {
                "Доля в %": share,
                "Доля в руб": clean_sum_text(summ),
                "Наимен. и реквизиты": full_name,
                "Дата": date
            }

            if any(record.values()):
                if has_strike or share.strip() == '–' or not share.strip():
                    outdated.append(record)
                else:
                    actual.append(record)

            i += 1

    return {
        "Актуальные участники": actual,
        "Неактуальные участники": outdated
    }


def extract_collaterals(doc):
    """Извлекает сведения о залогах"""
    collaterals = []
    date_pattern = re.compile(r'от\s+(\d{2}\.\d{2}\.\d{4})')

    for table in doc.tables:
        if len(table.columns) < 2:
            continue

        collateral_entry = {
            'Залогодатель': '',
            'Залогодержатель': '',
            'Дата залога': '',
            'Срок залога': '',
            'Заложенное имущество': ''
        }

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip().lower() if row.cells[0].text else ''
            value = row.cells[1].text.strip() if len(row.cells) > 1 and row.cells[1].text else ''

            if 'залогодатель' in key:
                collateral_entry['Залогодатель'] = value
            elif 'залогодержатель' in key:
                collateral_entry['Залогодержатель'] = value
            elif 'договор' in key:
                match = date_pattern.search(value)
                if match:
                    collateral_entry['Дата залога'] = match.group(1)
            elif 'срок исполн' in key:
                collateral_entry['Срок залога'] = value
            elif 'описание' in key:
                collateral_entry['Заложенное имущество'] = value

        if any(collateral_entry.values()):
            collaterals.append(collateral_entry)

    return collaterals


def extract_leasing_info(doc):
    """Извлекает информацию о лизинге"""
    leasing_info_all = []

    for table in doc.tables:
        if len(table.columns) != 2 or len(table.rows) < 5:
            continue

        first_cell_text = table.cell(0, 0).text.strip() if len(table.rows) > 0 and len(table.rows[0].cells) > 0 else ''
        if not re.match(r'\d{2}\.\d{2}\.\d{4}', first_cell_text):
            continue

        data = {
            "Лизингодатель": "",
            "Период лизинга": "",
            "Категория": "",
            "Текущий статус": ""
        }

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key_cell = row.cells[0].text.strip().lower() if row.cells[0].text else ''
            val_cell = row.cells[1].text.strip() if len(row.cells) > 1 and row.cells[1].text else ''

            if not key_cell:
                continue

            if key_cell.startswith("лизингодатель"):
                value = val_cell.lstrip("—–—-– ").strip()
                data["Лизингодатель"] = value

            elif key_cell.startswith("период лизинга"):
                data["Период лизинга"] = val_cell

            elif key_cell.startswith("категория"):
                data["Категория"] = val_cell

            elif key_cell.startswith("статус"):
                data["Текущий статус"] = val_cell.split('\n')[0].strip() if val_cell else ''

        if data["Лизингодатель"] or data["Категория"]:
            leasing_info_all.append(data)

    not_finished = [
        item for item in leasing_info_all
        if 'завершился' not in item.get('Текущий статус', '').lower()
    ]

    if not leasing_info_all:
        return ["Информации по лизингу нет"]
    if not not_finished:
        return ["Информации по лизингу со статусом 'Действует' нет"]
    return not_finished


def extract_credit_debt(doc):
    """Извлекает информацию о кредиторской задолженности по годам."""
    for table in doc.tables:
        if not table.rows or len(table.rows[0].cells) < 2:
            continue

        headers = []
        for cell in table.rows[0].cells:
            if cell.paragraphs and len(cell.paragraphs) > 0:
                headers.append(extract_text_without_strikethrough(cell.paragraphs[0]).lower())
            else:
                headers.append('')

        if (
                len(headers) >= 4
                and any('код' in h for h in headers)
                and sum(re.search(r'20\d{2}', h) is not None for h in headers) >= 2
        ):
            debt_row = None
            for row in table.rows:
                if len(row.cells) == 0:
                    continue
                first_cell = row.cells[0]
                if first_cell.paragraphs and len(first_cell.paragraphs) > 0:
                    text = extract_text_without_strikethrough(first_cell.paragraphs[0]).strip().lower()
                    if "кредиторская задолженность" in text:
                        debt_row = row
                        break

            if debt_row:
                year_val = {}
                for idx, head_text in enumerate(headers):
                    if idx >= len(debt_row.cells):
                        continue
                    year_match = re.search(r'(20\d{2})', head_text)
                    if year_match:
                        cell = debt_row.cells[idx]
                        if cell.paragraphs and len(cell.paragraphs) > 0:
                            value = extract_text_without_strikethrough(cell.paragraphs[0]).strip()
                            cleaned = value.replace(' ', '').replace('–', '0')
                            year_val[year_match.group(1)] = cleaned

                years_sorted = sorted(year_val.keys(), reverse=True)
                slots = ['year_1', 'year_2', 'year_3']
                res = {}
                for i, year in enumerate(years_sorted[:3]):
                    res[slots[i]] = {year: year_val[year]}
                return res

    return {'year_1': '', 'year_2': '', 'year_3': ''}


def extract_financial_results(doc):
    """Извлекает данные о выручке, чистой прибыли и убытке из таблиц."""
    revenue_data = {}

    for table in doc.tables:
        if not table.rows or len(table.rows[0].cells) < 2:
            continue

        headers = []
        for cell in table.rows[0].cells:
            if cell.paragraphs and len(cell.paragraphs) > 0:
                headers.append(extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower())
            else:
                headers.append('')

        if len(headers) < 2:
            continue

        has_revenue_row = False
        for row in table.rows:
            if len(row.cells) > 0 and row.cells[0].paragraphs:
                text = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower()
                if "выручка" in text:
                    has_revenue_row = True
                    break

        if not has_revenue_row:
            continue

        for row in table.rows:
            if not row.cells or len(row.cells) < 2:
                continue

            cell_0 = row.cells[0]
            if not (cell_0.paragraphs and len(cell_0.paragraphs) > 0):
                continue

            text = extract_text_without_strikethrough(cell_0.paragraphs[0]).strip().lower()
            if not any(key in text for key in ['выручка', 'чистая прибыль', 'убыток']):
                continue

            for idx, header in enumerate(headers):
                if idx >= len(row.cells):
                    continue

                year_match = re.search(r'(20\d{2})', header)
                if not year_match:
                    continue

                year = year_match.group(1)
                cell = row.cells[idx]
                if not (cell.paragraphs and len(cell.paragraphs) > 0):
                    continue

                value = extract_text_without_strikethrough(cell.paragraphs[0]).strip().replace(' ', '').replace('–',
                                                                                                                '0')
                revenue_data.setdefault(year, {})[text] = value

    years_sorted = sorted(revenue_data.keys(), reverse=True)
    slots = ['year_1', 'year_2', 'year_3']
    result = {}

    for i, year in enumerate(years_sorted[:3]):
        result[slots[i]] = {year: revenue_data[year]}

    return result if result else {'year_1': '', 'year_2': '', 'year_3': ''}


def extract_assets_and_receivables(doc):
    """Извлекает информацию об основных средствах и дебиторской задолженности."""
    result = {
        "Основные средства": {},
        "Дебиторская задолженность": {}
    }

    for table in doc.tables:
        header_cells = []
        if not table.rows or len(table.rows[0].cells) == 0:
            continue

        for cell in table.rows[0].cells:
            if cell.paragraphs and len(cell.paragraphs) > 0:
                header_cells.append(extract_text_without_strikethrough(cell.paragraphs[0]).strip().lower())
            else:
                header_cells.append('')
        year_to_col = {}

        for idx, txt in enumerate(header_cells):
            m = re.search(r'конец\s*(20\d{2})', txt)
            if m:
                year_to_col[m.group(1)] = idx

        if len(year_to_col) >= 3:
            years_sorted = sorted(year_to_col.keys(), reverse=True)[:3]
            row_ос = None
            row_дебит = None

            for row in table.rows:
                if len(row.cells) == 0:
                    continue

                first_cell = ''
                if row.cells[0].paragraphs and len(row.cells[0].paragraphs) > 0:
                    first_cell = extract_text_without_strikethrough(row.cells[0].paragraphs[0]).strip().lower()

                if first_cell == "основные средства":
                    row_ос = row
                if first_cell == "дебиторская задолженность":
                    row_дебит = row

            if row_ос and row_дебит:
                for i, year in enumerate(years_sorted):
                    idx = year_to_col[year]
                    val_os = ''
                    if idx < len(row_ос.cells) and row_ос.cells[idx].paragraphs:
                        val_os = extract_text_without_strikethrough(row_ос.cells[idx].paragraphs[0]).strip().replace(
                            ' ', '').replace('–', '0')
                    val_db = ''
                    if idx < len(row_дебит.cells) and row_дебит.cells[idx].paragraphs:
                        val_db = extract_text_without_strikethrough(row_дебит.cells[idx].paragraphs[0]).strip().replace(
                            ' ', '').replace('–', '0')

                    result["Основные средства"][f'year_{i + 1}'] = {year: val_os}
                    result["Дебиторская задолженность"][f'year_{i + 1}'] = {year: val_db}
                break

    return result


def extract_related_companies_from_path(filepath):
    """Ближайшие связи"""
    full_text = docx2txt.process(filepath)

    if not full_text or "Ближайшие связи – Актуальные" not in full_text:
        print("Блок 'Ближайшие связи – Актуальные' не найден.")
        return {"Ближайшие связи": {}}

    block_lines = full_text.split("Ближайшие связи – Актуальные", 1)[-1].splitlines()
    cleaned_block_lines = []
    for line in block_lines:
        if "Ближайшие связи –" in line and "Актуальные" not in line:
            break
        cleaned_block_lines.append(line.strip())

    lines = [l for l in cleaned_block_lines if l]

    related_companies = {}
    current = None
    current_key = None
    last_key_was_participant = False

    for i, line in enumerate(lines):
        clean_line = line.lstrip("▶►•").strip()

        if ("контур.фокус" in clean_line.lower()
                or "по данным" in clean_line.lower()
                or "на основании" in clean_line.lower()):
            continue

        if last_key_was_participant:
            last_key_was_participant = False
            continue

        if re.match(r'^(ООО|АО|ПАО|ИП)\s', clean_line):
            if re.search(r'ИНН.*\d{10,12}.*(%|руб\.|доля)', clean_line.lower()):
                continue
            if "%" in clean_line or "доля" in clean_line.lower() or "руб" in clean_line.lower():
                continue

            normalized_key = re.sub(r'\s+', '', clean_line).lower()
            if any(re.sub(r'\s+', '', k).lower() == normalized_key for k in related_companies):
                continue

            if current_key and current:
                related_companies[current_key] = current

            current_key = clean_line
            current = {
                "Адрес": "",
                "Генеральный директор": "" if current_key.startswith("ИП") else "",
                "Участники": [] if current_key.startswith("ИП") else [],
                "Реквизиты": {"ИНН": "", "ОГРН": ""}
            }
            continue

        if not current:
            continue

        if "инн" in clean_line.lower() or "огрн" in clean_line.lower():
            _, inn, ogrn = extract_inn_ogrn(clean_line)
            if inn:
                current["Реквизиты"]["ИНН"] = inn
            if ogrn:
                current["Реквизиты"]["ОГРН"] = ogrn

        if "Генеральный директор" in clean_line and i + 1 < len(lines):
            fio_line = lines[i + 1]
            cleaned, inn, _ = extract_inn_ogrn(fio_line)
            result = f"{cleaned}, ИНН {inn}" if inn else cleaned
            current["Генеральный директор"] = result

        if ("учредитель" in clean_line.lower() or "участник" in clean_line.lower()) and i + 1 < len(lines):
            full_text = lines[i + 1]
            participant = full_text.split("100%")[0].strip() + "100%" if "100%" in full_text else full_text
            participant = re.sub(r"\xa0", " ", participant)
            current["Участники"].append(participant)
            last_key_was_participant = True
            continue

        if clean_line.lower() == "адрес" and i + 1 < len(lines):
            current["Адрес"] = lines[i + 1]
        elif not current["Адрес"] and re.search(r'(г\.|г |обл\.|обл |респ\.|респ )', clean_line.lower()):
            current["Адрес"] = clean_line

    if current_key and current:
        related_companies[current_key] = current

    return related_companies


def extract_share_pledge_info(doc: Document):
    """Извлекает информацию о залоге долей с корректной обработкой ИНН."""

    def is_strikethrough(cell: _Cell) -> bool:
        """Проверяет, является ли текст в ячейке зачеркнутым."""
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.strike:
                    return True
        return False

    for table in doc.tables:
        if len(table.columns) != 4 or len(table.rows) < 2:
            continue

        rows = table.rows
        participants = []

        for i in range(1, len(rows)):
            row = rows[i]
            if len(row.cells) < 4:
                continue

            name_cell = row.cells[2]
            date_cell = row.cells[3]

            name_text = extract_text_from_cell(name_cell)
            date_text = extract_text_from_cell(date_cell)

            if not name_text or is_strikethrough(name_cell):
                continue

            if "Залог доли Залогодержатель" not in name_text:
                participants.append({
                    "index": i,
                    "name": name_text
                })

        for participant in participants:
            next_idx = participant["index"] + 1
            if next_idx < len(rows):
                next_row = rows[next_idx]
                if len(next_row.cells) < 4:
                    continue

                pledge_cell = next_row.cells[2]
                pledge_date_cell = next_row.cells[3]

                pledge_text = extract_text_from_cell(pledge_cell)
                pledge_date = extract_text_from_cell(pledge_date_cell)

                if "Залог доли Залогодержатель" in pledge_text and not is_strikethrough(pledge_cell):
                    if ":" in pledge_text:
                        pledge_holder_raw = pledge_text.split(":", 1)[1].strip()
                    else:
                        pledge_holder_raw = pledge_text

                    pledge_holder_clean, inn, _ = extract_inn_ogrn(pledge_holder_raw)
                    pledge_holder = f"{pledge_holder_clean}, ИНН {inn}" if inn else pledge_holder_clean

                    return {
                        "Залогодатель": participant["name"],
                        "Залогодержатель": pledge_holder,
                        "Дата залога": pledge_date
                    }

    return "Долей в залоге нет"


def log_result(name: str, result):
    success = False
    if isinstance(result, str):
        success = bool(result.strip())
    elif isinstance(result, dict):
        success = any(bool(v) for v in result.values())
    elif isinstance(result, list):
        success = len(result) > 0
    print(f"{name}: {'✅' if success else '❌'}")


def parsing_all_docx(docx_path: str) -> dict:
    doc = Document(docx_path)
    company_data = {}

    print("\nПАРСИНГ .docx ФАЙЛА:")

    # 1. Собираем все данные с логированием
    basic_info = extract_basic_info(doc)
    log_result("Общая информация", basic_info)

    founders = extract_founders(doc)
    employees = extract_staff_info(doc)
    log_result("Сведения о сотрудниках", employees)

    share_pledge_info = extract_share_pledge_info(doc)
    log_result("Залог долей", share_pledge_info)

    connections = extract_related_companies_from_path(docx_path)
    log_result("Ближайшие связи", connections)

    assets = extract_assets_and_receivables(doc)
    log_result("Основные средства и дебиторка", assets)

    collaterals = extract_collaterals(doc)
    log_result("Сведения о залогах", collaterals)

    leasing = extract_leasing_info(doc)
    log_result("Сведения о лизинге", leasing)

    credit_debt = extract_credit_debt(doc)
    log_result("Кредиторская задолженность", credit_debt)

    financials = extract_financial_results(doc)
    log_result("Финансовые результаты", financials)

    log_result("Учредители/участники", founders)

    # 2. Формируем раздел "Общая информация" с добавлением конкурсного управляющего
    manager = extract_competitive_manager(doc)
    general_info = {
        "Общая информация": {
            "Наименование": basic_info.get("Наименование", ""),
            "ОГРН": basic_info.get("ОГРН", ""),
            "ИНН/КПП": f"{basic_info.get('ИНН', '')} / {basic_info.get('КПП', '')}",
            "Юридический адрес": basic_info.get("Юридический адрес", ""),
            "Дата создания": basic_info.get("Дата создания", ""),
            "Учредители/участники (текущие)": format_founders(founders.get("Актуальные участники", [])),
            "Размер уставного капитала": basic_info.get("Уставный капитал", ""),
            "Генеральный директор": f"{basic_info.get('Генеральный директор', {}).get('ФИО', '')}, ИНН {basic_info.get('Генеральный директор', {}).get('ИНН', '')}".strip(
                ", "),
            "Конкурсный управляющий": f"{manager.get('ФИО', '')}, ИНН {manager.get('ИНН', '')}".strip(", "),
            "ОКВЭД (основной)": basic_info.get("ОКВЭД (основной)", "")
        }
    }

    # 3. Формируем результат в нужном порядке
    result = {
        'Общая информация': general_info["Общая информация"],
        'Сведения о сотрудниках': employees,
        'Залог долей': share_pledge_info,
        'Ближайшие связи': connections,
        'Основные средства и дебиторка': assets,
        'Сведения о залогах': collaterals,
        'Сведения о лизинге': leasing,
        'Кредиторская задолженность': credit_debt,
        'Финансовые результаты': financials,
        'Учредители/участники': founders
    }

    # 4. Выводим результат в требуемом формате
    print("\nРЕЗУЛЬТАТ ПАРСИНГА .docx ФАЙЛА:")
    pprint(result)

    return result


if __name__ == "__main__":
    docx_path = os.path.join(os.path.dirname(__file__), "word.docx")
    result = parsing_all_docx(docx_path)
    from pprint import pprint

    pprint(result)