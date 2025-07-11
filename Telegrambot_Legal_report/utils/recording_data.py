# Занесение в шаблон
import os
import re
from docx import Document
import ast
from datetime import datetime
from docx.shared import RGBColor


def generate_filename(data: dict) -> str:
    """
    Генерирует имя файла на основе наименования организации и текущего времени.
    Формат: <Наименование>_<ГГГГ-ММ-ДД_ЧЧММСС>.docx
    """
    info = data.get("Общая информация", {})
    name = info.get("Наименование") or "Без названия"
    name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '', name)
    name = name.replace('"', '').replace("«", "").replace("»", "").strip()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    return f"{name}_{timestamp}.docx"


def safe_str(val):
    """Преобразует значение в безопасную строку, удаляя непечатаемые символы типа u0000."""
    if val is None:
        return ""
    s = str(val)
    return ''.join(c for c in s if c.isprintable())


def fill_table1(table, data: dict):
    """Заполнение таблицы 1 - Основные сведения о компании"""
    info = data.get("Общая информация", {})

    # Обработка учредителей
    actual_founders = data.get('Учредители/участники', {}).get('Актуальные участники', [])
    founders_text = ""
    for i, founder in enumerate(actual_founders, 1):
        name = founder.get('Наимен. и реквизиты', '')
        share = founder.get('Доля в %', '')
        founders_text += f"{i}. {name} — {share}\n"
    founders_text = founders_text.strip()

    # Разбор ИНН/КПП
    inn_kpp = info.get("ИНН/КПП", "")
    if " / " in inn_kpp:
        inn, kpp = inn_kpp.split(" / ")
    else:
        inn, kpp = inn_kpp, ""

    # Генеральный директор
    director_raw = info.get("Генеральный директор", "")
    fio = ""
    dir_inn = ""
    if "ИНН" in director_raw:
        parts = director_raw.split("ИНН")
        fio = parts[0].strip(", ")
        dir_inn = parts[1].strip()

    # Заполнение ячеек
    # Наименование — белым цветом
    cell = table.cell(0, 1)
    cell.text = ""
    run = cell.paragraphs[0].add_run(info.get("Наименование", ""))
    run.font.color.rgb = RGBColor(255, 255, 255)
    table.cell(1, 1).text = info.get("ОГРН", "")
    table.cell(2, 1).text = f"{inn} / {kpp}"
    table.cell(3, 1).text = info.get("Юридический адрес", "")
    table.cell(4, 1).text = info.get("Дата создания", "")
    table.cell(5, 1).text = founders_text
    table.cell(6, 1).text = info.get("Размер уставного капитала", "")
    table.cell(7, 0).text = "Генеральный директор"
    table.cell(7, 1).text = f"{fio}, ИНН {dir_inn}".strip(", ")
    table.cell(8, 1).text = info.get("ОКВЭД (основной)", "")
    table.cell(9, 1).text = info.get("Система налогообложения", "")


from docx.shared import RGBColor

def fill_table2(table, data: dict):
    """Заполнение таблицы 2 - Сведения о сотрудниках с белыми надписями"""
    def extract_value(dictionary_str):
        try:
            parsed = ast.literal_eval(dictionary_str)
            if isinstance(parsed, dict):
                key = next(iter(parsed.keys()))
                val = parsed[key]
                return key, val
        except Exception:
            return "", ""
        return "", ""

    employee_info = data.get("Сведения о сотрудниках", {})
    headcount = employee_info.get("Среднесписочная численность", {})
    salary = employee_info.get("Средняя заработная плата", {})

    columns = ['year_3', 'year_2', 'year_1']
    years, headcount_values, salary_values = [], [], []

    for year_key in columns:
        y, v1 = extract_value(headcount.get(year_key, ''))
        _, v2 = extract_value(salary.get(year_key, ''))
        years.append(y)
        headcount_values.append(v1)
        salary_values.append(v2)

    # Заголовки лет — белым
    for i, year in enumerate(years):
        cell = table.cell(0, i + 1)
        cell.text = ""
        run = cell.paragraphs[0].add_run(year)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Значения численности и зарплат — обычные
    for i, val in enumerate(headcount_values):
        table.cell(1, i + 1).text = val
    for i, val in enumerate(salary_values):
        table.cell(2, i + 1).text = val

    # Подписи строк слева — белым
    for row_idx, label in [(1, "Среднесписочная численность"), (2, "Средняя заработная плата")]:
        cell = table.cell(row_idx, 0)
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.font.color.rgb = RGBColor(255, 255, 255)


# def extract_dates_and_names(data: dict):
#     """
#     Извлекает все уникальные даты и соответствующие доли из участников.
#     """
#     import re
#
#     def normalize_date(text):
#         match = re.search(r"\d{2}\.\d{2}\.\d{4}", text)
#         return match.group(0) if match else None
#
#     date_set = set()
#     rows = []
#
#     for founder in data.get("Актуальные участники", []) + data.get("Неактуальные участники", []):
#         name = founder.get("Наимен. и реквизиты", "")
#         share = founder.get("Доля в %", "")
#         date = founder.get("Дата", "")
#         date_points = []
#
#         if "(" in date:
#             date_parts = [normalize_date(d) for d in date.split("(")]
#             date_points = [d for d in date_parts if d]
#         else:
#             d = normalize_date(date)
#             if d:
#                 date_points = [d]
#
#         for d in date_points:
#             date_set.add(d)
#
#         is_old = founder in data.get("Неактуальные участники", [])
#         rows.append((name, share, date_points, is_old))
#
#     sorted_dates = sorted(date_set, key=lambda d: datetime.strptime(d, "%d.%m.%Y"))
#     return sorted_dates, rows


# def insert_timeline_share_table(document: Document, data: dict) -> bool:
#     """
#     Вставляет таблицу с временной шкалой владения долями после нужного заголовка.
#     """
#     sorted_dates, founder_rows = extract_dates_and_names(data)
#
#     # Найти нужный параграф
#     target_paragraph = None
#     for paragraph in document.paragraphs:
#         if "Хронология владения долями в уставном капитале" in paragraph.text:
#             target_paragraph = paragraph
#             break
#     if not target_paragraph:
#         print("❌ Не найден заголовок 'Хронология владения долями'")
#         return False
#
#     # Создаём таблицу
#     table = document.add_table(rows=1, cols=1 + len(sorted_dates))
#     table.style = "Table Grid"
#     table.alignment = WD_TABLE_ALIGNMENT.LEFT
#
#     hdr = table.rows[0].cells
#     hdr[0].text = "Наименование"
#     for i, date in enumerate(sorted_dates):
#         hdr[i + 1].text = date
#
#     for name, share, date_points, is_old in founder_rows:
#         row = table.add_row().cells
#         row[0].text = name
#         for d in date_points:
#             try:
#                 col_idx = sorted_dates.index(d)
#                 row[col_idx + 1].text = share
#             except ValueError:
#                 pass
#         if is_old:
#             for cell in row:
#                 for p in cell.paragraphs:
#                     for r in p.runs:
#                         r.font.strike = True
#
#     # Вставляем таблицу после заголовка
#     target_paragraph._p.addnext(table._tbl)
#     return True


def fill_table4(table, data: dict):
    """Заполнение таблицы 4 — Сведения о залоге долей, или удаление таблицы"""
    value = data.get("Залог долей")

    # if isinstance(value, str) and "нет" in value.lower():
    #     # Удаляем таблицу, если информация отсутствует
    #     tbl_element = table._tbl
    #     parent_element = tbl_element.getparent()
    #     parent_element.remove(tbl_element)
    #
    #     # Добавляем параграф вместо таблицы
    #     paragraph = table._parent.add_paragraph("Залог долей нет")
    #     run = paragraph.runs[0]
    #     run.italic = True
    #     return
    # if not isinstance(value, dict):
    #     return

    if not value or (isinstance(value, str) and "нет" in value.lower()):
        return

    table.cell(1, 0).text = value.get("Залогодатель", "")
    table.cell(1, 1).text = value.get("Дата залога", "")
    table.cell(1, 2).text = value.get("Залогодержатель", "")


def fill_table5(table, data: dict):
    """Заполнение таблицы 5 - Аффилированность и ближайшие связи"""
    # Удаляем все строки после заголовка (оставляем только строку с заголовками)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # Получаем данные
    links_data = data.get("Ближайшие связи", {})
    if isinstance(links_data, dict) and "Ближайшие связи" in links_data:
        links_data = links_data["Ближайшие связи"]

    for company_name, company_info in links_data.items():
        row_cells = table.add_row().cells

        # 1. Наименование и ИНН
        inn = company_info.get("Реквизиты", {}).get("ИНН", "")
        row_cells[0].text = f"{company_name}\nИНН: {inn}"

        # 2. Генеральный директор
        gen_dir = company_info.get("Генеральный директор", "").strip(", ")
        row_cells[1].text = gen_dir

        # 3. Участники
        participants = company_info.get("Участники", [])
        formatted_participants = "\n".join(
            f"{i + 1}) {p.strip('► ').strip()}" for i, p in enumerate(participants)
        )
        row_cells[2].text = formatted_participants

        # 4. Адрес
        row_cells[3].text = company_info.get("Адрес", "")

        # 5. Взаимосвязь — пусто или логика
        row_cells[4].text = ""


def fill_table6(table, data: dict):
    """Заполнение таблицы 6 — Сведения о размере основных средств и дебиторской задолженности"""
    values = data.get("Основные средства и дебиторка", {})
    fixed_assets = values.get("Основные средства", {})
    receivables = values.get("Дебиторская задолженность", {})

    # Получаем список годов и значений
    year_map = {}  # year: column_index

    years = []
    for key in ['year_3', 'year_2', 'year_1']:
        y_fixed = fixed_assets.get(key, {})
        y_receiv = receivables.get(key, {})

        year = next(iter(y_fixed or y_receiv), None)
        if year:
            years.append(year)

    # Заполняем заголовки годов — белым цветом
    for col_idx, year in enumerate(years):
        if col_idx + 1 < len(table.columns):
            cell = table.cell(0, col_idx + 1)
            cell.text = ""
            run = cell.paragraphs[0].add_run(year)
            run.font.color.rgb = RGBColor(255, 255, 255)
            year_map[year] = col_idx + 1

    # Заполнение строки "Размер основных средств"
    for key, year_data in fixed_assets.items():
        for year, val in year_data.items():
            col = year_map.get(year)
            if col:
                table.cell(1, col).text = f"{int(val):,}".replace(",", " ") + " руб."

    # Заполнение строки "Дебиторская задолженность"
    for key, year_data in receivables.items():
        for year, val in year_data.items():
            col = year_map.get(year)
            if col:
                table.cell(2, col).text = f"{int(val):,}".replace(",", " ") + " руб."

def fill_table8(table, data: dict):
    """Заполнение таблицы 8 — Сведения о залогах (Залогодержатель, дата, срок, имущество)"""
    items = data.get("Сведения о залогах", [])

    # Удаляем все строки, кроме заголовка
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for idx, item in enumerate(items, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item.get("Залогодержатель", "")
        row[2].text = item.get("Дата залога", "")
        row[3].text = item.get("Срок залога", "") or ""
        row[4].text = item.get("Заложенное имущество", "") or ""


def fill_table9(table, data: dict):
    """Заполнение таблицы 9 — Сведения о лизинге с подстановкой если лизингодатель отсутствует"""
    leasers = data.get("Сведения о лизинге", [])

    if not isinstance(leasers, list):
        print("ОШИБКА: Сведения о лизинге не являются списком.")
        return

    # Удаляем все строки кроме заголовка
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for idx, lease in enumerate(leasers, start=1):
        if not isinstance(lease, dict):
            continue

        row_cells = table.add_row().cells

        row_cells[0].text = str(idx)

        # Если нет значения — подставляем формулировку
        leaser = lease.get("Лизингодатель", "").strip()
        row_cells[1].text = leaser if leaser else "Сведения о лизингодателе отсутствуют"

        row_cells[2].text = lease.get("Период лизинга", "")
        row_cells[3].text = lease.get("Категория", "")
        row_cells[4].text = lease.get("Текущий статус", "")


def fill_table10(table, data):
    """
     Заполняет таблицу 'Сведения о просуживаемой задолженности'.
     """
    items = data.get("Просуживаемая", [])
    if not isinstance(items, list):
        print("⚠️ Ожидается список в поле 'Просуживаемая'")
        return

    # Удаляем все строки, кроме заголовка
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for item in items:
        row = table.add_row().cells
        row[0].text = item.get("№ Дела", "")
        row[1].text = item.get("Ответчик", "")
        row[2].text = item.get("Исковые требования", "")
        row[3].text = item.get("Статус", "")


def fill_table11(table, data):
    """
    Заполняет таблицу 'Сведения о размере кредиторской задолженности по бух. балансу'
    с белыми заголовками годов.
    """
    if not isinstance(data, dict):
        raise ValueError("Ожидается словарь для 'Кредиторская задолженность'")

    year_data = {}
    for key in ['year_1', 'year_2', 'year_3']:
        year_data.update(data.get(key, {}))

    if not year_data:
        raise ValueError("Нет данных для заполнения кредиторской задолженности")

    sorted_years = sorted(year_data.keys())[:3]

    if len(table.rows) < 2 or len(table.columns) < 4:
        raise ValueError("Ожидается таблица минимум 2 строки и 4 столбца")

    for col_idx, year in enumerate(sorted_years, start=1):
        try:
            # Белый заголовок года
            cell = table.cell(0, col_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(year))
            run.font.color.rgb = RGBColor(255, 255, 255)

            # Значение
            value = year_data[year]
            value_int = int(str(value).replace(" ", "").replace(",", ""))
            table.cell(1, col_idx).text = f"{value_int:,} т.р.".replace(",", " ")
        except Exception as e:
            print(f"Ошибка при вставке значения для {year}: {e}")
            raise


def fill_table12(table, data: dict):
    """
    Заполняет таблицу 'Сведения о просуженной кредиторской задолженности'.
    """
    items = data.get("Просуженная", [])
    if not isinstance(items, list):
        print("⚠️ Ожидается список в поле 'Просуженная'")
        return

    # Удаляем все строки кроме заголовка
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for item in items:
        row = table.add_row().cells
        row[0].text = item.get("№ Дела", "")
        row[1].text = item.get("Ответчик", "")
        row[2].text = item.get("Исковые требования", "")
        row[3].text = item.get("Статус", "")



def fill_table13(table, data):
    """
    Заполняет таблицу 'Отчет о финансовых результатах' с белыми заголовками годов.
    """
    if not isinstance(data, dict):
        print("❌ Неверный формат данных для финансовых результатов")
        return

    # Собираем все года
    all_years = set()
    for values in data.values():
        if isinstance(values, dict):
            all_years.update(values.keys())

    if len(all_years) != 4:
        print(f"❌ Ожидается ровно 4 года, получено: {sorted(all_years)}")
        return

    years = sorted(all_years)
    col_by_year = {year: col_idx + 1 for col_idx, year in enumerate(years)}

    # Белые заголовки годов в первой строке
    for year, col_idx in col_by_year.items():
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(f"конец {year}")
        run.font.color.rgb = RGBColor(255, 255, 255)

    filled, skipped = 0, 0

    # Заполнение значений
    for row in table.rows[1:]:
        indicator = row.cells[0].text.strip()
        if not indicator or indicator not in data:
            continue

        year_values = data[indicator]
        for year, value in year_values.items():
            col_idx = col_by_year.get(year)
            if col_idx is None:
                skipped += 1
                continue
            try:
                row.cells[col_idx].text = f"{int(value):,}".replace(",", " ")
                filled += 1
            except Exception as e:
                print(f"❌ Ошибка при вставке '{indicator}' за {year}: {e}")
                skipped += 1


def fill_table14(table, data: dict):
    """
    Заполняет таблицу 'Финансовый анализ на последнюю отчетную дату' с белыми заголовками годов.
    """
    analysis = data.get("Финансовый анализ", {})
    if not isinstance(analysis, dict):
        print("⚠️ Неверный формат данных для Финансового анализа.")
        return

    # Извлекаем года
    years = set()
    for val in analysis.values():
        if isinstance(val, dict):
            years.update(k for k in val if k.startswith("Значение показателя на "))

    if len(years) != 2:
        print(f"❌ Ожидается ровно 2 года, найдено: {sorted(years)}")
        return

    sorted_years = sorted(years)
    year_col_map = {year: i + 1 for i, year in enumerate(sorted_years)}  # колонки: 1, 2

    # Заголовки годов — белым цветом
    for year, col in year_col_map.items():
        cell = table.cell(0, col)
        cell.text = ""
        run = cell.paragraphs[0].add_run(year)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Заголовок последней колонки (описание) — оставляем как есть

    # Заполнение значений
    for row in table.rows[1:]:
        indicator = row.cells[0].text.strip()
        if not indicator:
            continue

        row_data = analysis.get(indicator)
        if not isinstance(row_data, dict):
            continue

        for year, col in year_col_map.items():
            value = row_data.get(year)
            if value is not None:
                row.cells[col].text = str(value)

        desc = row_data.get("Описание показателя")
        if desc:
            row.cells[3].text = desc


def save_filled_doc(template_path: str, output_path: str, data: dict) -> str:
    """Заполнение шаблона Word-отчета по данным и сохранение файла с именем по наименованию и дате"""
    document = Document(template_path)
    status = {}

    try:
        fill_table1(document.tables[0], data)
        status["Общая информация"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 1 (Общая информация):", e)
        status["Общая информация"] = False

    try:
        fill_table2(document.tables[1], data)
        status["Сведения о сотрудниках"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 2 (Сведения о сотрудниках):", e)
        status["Сведения о сотрудниках"] = False

    # try:
    #     insert_timeline_share_table(document, data.get("Учредители/участники", {}))
    #     status["Хронология владения долями"] = True
    # except Exception as e:
    #     print("Ошибка при добавлении таблицы хронологии долей:", e)
    #     status["Хронология владения долями"] = False

    try:
        fill_table4(document.tables[3], data)
        status["Залог долей"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 4 (Залог долей):", e)
        status["Залог долей"] = False

    try:
        fill_table5(document.tables[4], data)
        status["Ближайшие связи"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 5 (Ближайшие связи):", e)
        status["Ближайшие связи"] = False

    try:
        fill_table6(document.tables[5], data)
        status["Основные средства и дебиторка"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 6 (Основные средства и дебиторка):", e)
        status["Основные средства и дебиторка"] = False

    try:
        fill_table8(document.tables[7], data)
        status["Сведения о залогах"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 8 (Сведения о залогах):", e)
        status["Сведения о залогах"] = False

    try:
        fill_table9(document.tables[8], data)
        status["Сведения о лизинге"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 9 (Сведения о лизинге):", e)
        status["Сведения о лизинге"] = False

    try:
        fill_table10(document.tables[9], data)
        status["Просуживаемая задолженность"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 9 (Просуживаемая задолженность):", e)
        status["Просуживаемая задолженность"] = False

    try:
        fill_table11(document.tables[10], data["Кредиторская задолженность"])
        status["Кредиторская задолженность"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 10 (Кредиторская задолженность):", e)
        status["Кредиторская задолженность"] = False

    try:
        fill_table12(document.tables[11], data)
        status["Просуженная задолженность"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 11 (Просуженная задолженность):", e)
        status["Просуженная задолженность"] = False

    try:
        fill_table13(document.tables[12], data["Финансовые результаты"])
        status["Финансовые результаты"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 13 (Финансовые результаты):", e)
        status["Финансовые результаты"] = False

    try:
        fill_table14(document.tables[13], data)
        status["Финансовый анализ"] = True
    except Exception as e:
        print("Ошибка при заполнении таблицы 14 (Финансовый анализ):", e)
        status["Финансовый анализ"] = False

    # Сохранение по пути output_path (переопределим имя на основе данных)
    filename = generate_filename(data)
    final_path = os.path.join("Reports", filename)
    document.save(final_path)

    # Вывод финального отчета
    print("\nЗАПОЛНЕНИЕ ТАБЛИЦ:")
    for section in [
        "Общая информация",
        "Сведения о сотрудниках",
        "Залог долей",
        "Ближайшие связи",
        "Основные средства и дебиторка",
        "Сведения о залогах",
        "Сведения о лизинге",
        "Просуживаемая задолженность",
        "Кредиторская задолженность",
        "Просуженная задолженность",
        "Финансовые результаты",
        "Финансовый анализ",
        "Общая информация",
        "Сведения о сотрудниках",
        "Хронология владения долями",
    ]:
        mark = "✅" if status.get(section, False) else "❌"
        print(f"{section}: {mark}")

    return final_path