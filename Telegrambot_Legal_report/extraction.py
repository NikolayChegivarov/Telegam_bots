from docx import Document
import re


def get_doc_text(file_path: str) -> str:
    """Извлекает весь текст из документа (параграфы + таблицы)."""
    doc = Document(file_path)
    full_text = []

    # Текст из параграфов
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())

    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    return '\n'.join(full_text)


def extract_org_data(file_path: str) -> dict:
    """Извлекает наименование организации, ИНН, КПП и ОГРН из документа. Возвращает словарь."""
    doc_text = get_doc_text(file_path)

    # Инициализируем словарь
    result = {'Организация': None, 'ОГРН': None, 'ИНН': None, 'КПП': None}

    # Ищем наименование организации (первая строка с организационной формой и кавычками)
    org_match = re.search(r'^((?:ООО|АО|ЗАО|ИП|ПАО|ОАО)\s*["«][^"»]+["»])', doc_text)

    if org_match:
        result['Организация'] = org_match.group(1).strip()
        print(f"Наименование организации нашли: {result['Организация']}")
    else:
        print("Наименование организации не нашли")
        print("Начало текста документа:", doc_text[:200])

    # Ищем ИНН (10 или 12 цифр)
    inn_match = re.search(r'ИНН[\s:–-]*(\d{10,12})', doc_text, re.IGNORECASE)
    if inn_match:
        result['ИНН'] = inn_match.group(1)
        print(f"ИНН нашли: {result['ИНН']}")
    else:
        print("ИНН не нашли")

    # Ищем КПП (9 цифр)
    kpp_match = re.search(r'КПП[\s:–-]*(\d{9})', doc_text, re.IGNORECASE)
    if kpp_match:
        result['КПП'] = kpp_match.group(1)
        print(f"КПП нашли: {result['КПП']}")
    else:
        print("КПП не нашли")

    # Ищем ОГРН (13 цифр)
    ogrn_match = re.search(r'ОГРН[\s:–-]*(\d{13})', doc_text, re.IGNORECASE)
    if ogrn_match:
        result['ОГРН'] = ogrn_match.group(1)
        print(f"ОГРН нашли: {result['ОГРН']}")
    else:
        print("ОГРН не нашли")

    return result


# Пример использования
if __name__ == "__main__":
    file_path = 'контур_фокус.docx'
    org_data = extract_org_data(file_path)

    print("Результат:")
    print(org_data)
