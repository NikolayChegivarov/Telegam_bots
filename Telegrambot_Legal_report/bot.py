import os
import json
from datetime import datetime
from dotenv import load_dotenv
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler,
)

# Загрузка переменных окружения
load_dotenv()
TOKEN = os.getenv('TELEGRAM_TOKEN_BOT')

print(f"Загружен токен бота: {'***' if TOKEN else 'НЕ НАЙДЕН!'}")

# Состояния для ConversationHandler
WAITING_FILE, WAITING_REPORT_NAME = range(2)


# Клавиатура с основными кнопками
def get_main_keyboard():
    keyboard = [
        [KeyboardButton("Создать отчет")],
        [KeyboardButton("История запросов"), KeyboardButton("Извлечь отчет")]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)


# Функции из ваших файлов (адаптированные)
def create_folder():
    """Создает папку Reports в корневой директории, если её нет"""
    reports_dir = "Reports"
    print(f"Проверка существования папки: {reports_dir}")
    if not os.path.exists(reports_dir):
        try:
            os.makedirs(reports_dir)
            print(f"✔ Создана папка: {reports_dir}")
        except Exception as e:
            print(f"❌ Ошибка при создании папки {reports_dir}: {str(e)}")
            return None
    return reports_dir


def create_history_file():
    """Создает файл history_requests.json в корневой директории, если его нет"""
    history_file = "history_requests.json"
    print(f"Проверка существования файла истории: {history_file}")
    if not os.path.exists(history_file):
        try:
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump({}, f)
            print(f"✔ Создан файл истории запросов: {history_file}")
        except Exception as e:
            print(f"❌ Ошибка при создании файла истории: {str(e)}")
    return history_file


def write_to_history(org_name):
    """Записывает новую запись в файл истории"""
    history_file = "history_requests.json"
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"Запись в историю: {org_name} - {current_date}")

    try:
        with open(history_file, 'r', encoding='utf-8') as f:
            history_data = json.load(f)

        history_data[org_name] = current_date

        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history_data, f, ensure_ascii=False, indent=4)

        print(f"✔ Добавлена запись в историю: {org_name} - {current_date}")
    except Exception as e:
        print(f"❌ Ошибка при записи в историю: {str(e)}")


def extract_org_data(file_path: str) -> dict:
    """Извлекает наименование организации, ИНН, КПП и ОГРН из документа."""
    from docx import Document
    import re

    print(f"Извлечение данных из файла: {file_path}")

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"❌ Ошибка при открытии файла DOCX: {str(e)}")
        return {'Организация': None, 'ОГРН': None, 'ИНН': None, 'КПП': None}

    full_text = []

    # Текст из параграфов
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())

    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    doc_text = '\n'.join(full_text)
    print(f"Извлеченный текст (первые 200 символов): {doc_text[:200]}...")

    # Инициализируем словарь
    result = {'Организация': None, 'ОГРН': None, 'ИНН': None, 'КПП': None}

    # Ищем наименование организации
    org_match = re.search(r'^((?:ООО|АО|ЗАО|ИП|ПАО|ОАО)\s*["«][^"»]+["»])', doc_text)
    if org_match:
        result['Организация'] = org_match.group(1).strip()
        print(f"Найдено наименование организации: {result['Организация']}")
    else:
        print("❌ Наименование организации не найдено")

    # Ищем ИНН (10 или 12 цифр)
    inn_match = re.search(r'ИНН[\s:–-]*(\d{10,12})', doc_text, re.IGNORECASE)
    if inn_match:
        result['ИНН'] = inn_match.group(1)
        print(f"Найден ИНН: {result['ИНН']}")
    else:
        print("❌ ИНН не найден")

    # Ищем КПП (9 цифр)
    kpp_match = re.search(r'КПП[\s:–-]*(\d{9})', doc_text, re.IGNORECASE)
    if kpp_match:
        result['КПП'] = kpp_match.group(1)
        print(f"Найден КПП: {result['КПП']}")
    else:
        print("❌ КПП не найден")

    # Ищем ОГРН (13 цифр)
    ogrn_match = re.search(r'ОГРН[\s:–-]*(\d{13})', doc_text, re.IGNORECASE)
    if ogrn_match:
        result['ОГРН'] = ogrn_match.group(1)
        print(f"Найден ОГРН: {result['ОГРН']}")
    else:
        print("❌ ОГРН не найден")

    return result


def process_template(template_path, org_data):
    """Обрабатывает шаблон и возвращает путь к сохраненному файлу"""
    from docx import Document

    print(f"\nОбработка шаблона: {template_path}")
    print("Полученные данные:", org_data)

    try:
        doc = Document(template_path)
        print("✔ Шаблон успешно загружен")
    except Exception as e:
        print(f"❌ Ошибка при открытии шаблона: {str(e)}")
        return None

    if not doc.tables:
        print("❌ В документе нет таблиц!")
        return None

    table = doc.tables[0]
    print(f"Таблица найдена: {len(table.rows)} строк, {len(table.columns)} столбцов")

    try:
        # Заполняем данные в таблице
        name_cell = table.cell(0, 1)
        org_name = org_data.get('Организация', 'Наименование не найдено')
        name_cell.text = org_name
        print(f"Заполнено наименование: {org_name}")

        ogrn_cell = table.cell(1, 1)
        ogrn = org_data.get('ОГРН', 'ОГРН не найден')
        ogrn_cell.text = ogrn
        print(f"Заполнен ОГРН: {ogrn}")

        inn_kpp_cell = table.cell(2, 1)
        inn = org_data.get('ИНН', 'ИНН не найден')
        kpp = org_data.get('КПП', 'КПП не найден')
        inn_kpp_cell.text = f"{inn}/{kpp}"
        print(f"Заполнен ИНН/КПП: {inn}/{kpp}")

        # Создаем папку Reports
        reports_dir = create_folder()
        if not reports_dir:
            return None

        # Формируем имя файла
        clean_name = "".join(c for c in org_name if c.isalnum() or c in (' ', '_', '-'))
        clean_name = clean_name.strip()[:50] or "Отчет_без_названия"
        output_filename = f"{clean_name}.docx"
        full_output_path = os.path.join(reports_dir, output_filename)
        print(f"Формируем имя файла: {full_output_path}")

        # Сохраняем файл
        doc.save(full_output_path)
        print(f"✔ Файл успешно сохранен: {full_output_path}")

        # Записываем в историю
        if org_name != 'Наименование не найдено':
            write_to_history(org_name)

        return full_output_path
    except Exception as e:
        print(f"❌ Ошибка при обработке шаблона: {str(e)}")
        return None


# Обработчики команд бота
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print(f"\nПолучена команда /start от пользователя {update.effective_user.id}")
    await update.message.reply_text(
        "Добро пожаловать! Выберите действие:",
        reply_markup=get_main_keyboard()
    )


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    user = update.effective_user
    print(f"\nПолучено сообщение от {user.id}: {text}")

    if text == "Создать отчет":
        print("Пользователь выбрал 'Создать отчет'")
        await update.message.reply_text(
            "Пожалуйста, прикрепите файл с юридической информацией.",
            reply_markup=get_main_keyboard()
        )
        return WAITING_FILE
    elif text == "История запросов":
        print("Пользователь выбрал 'История запросов'")
        await show_history(update, context)
    elif text == "Извлечь отчет":
        print("Пользователь выбрал 'Извлечь отчет'")
        await update.message.reply_text(
            "Введите название отчета (без расширения .docx):",
            reply_markup=get_main_keyboard()
        )
        return WAITING_REPORT_NAME
    else:
        print("Получено неизвестное текстовое сообщение")
        await update.message.reply_text(
            "Пожалуйста, используйте кнопки для навигации.",
            reply_markup=get_main_keyboard()
        )


async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    print(f"\nОбработка файла от пользователя {user.id}")

    # Проверяем, что получен документ
    if not update.message.document:
        print("❌ Получено сообщение без документа")
        await update.message.reply_text("Пожалуйста, прикрепите файл в формате DOCX.")
        return WAITING_FILE

    print(f"Получен документ: {update.message.document.file_name}")

    # Скачиваем файл
    file = await context.bot.get_file(update.message.document.file_id)
    temp_file = "temp.docx"
    print(f"Скачивание файла во временный файл: {temp_file}")
    await file.download_to_drive(temp_file)

    # Извлекаем данные
    org_data = extract_org_data(temp_file)

    # Обрабатываем шаблон
    template_path = "шаблон.docx"
    if not os.path.exists(template_path):
        print("❌ Файл шаблона не найден")
        await update.message.reply_text("Ошибка: файл шаблона не найден.")
        os.remove(temp_file)
        return ConversationHandler.END

    result_path = process_template(template_path, org_data)
    os.remove(temp_file)
    print(f"Временный файл {temp_file} удален")

    if result_path:
        print(f"Отправка результата пользователю: {result_path}")
        with open(result_path, 'rb') as f:
            await update.message.reply_document(
                document=f,
                caption="Ваш отчет готов!",
                reply_markup=get_main_keyboard()
            )
    else:
        print("❌ Ошибка при создании отчета")
        await update.message.reply_text(
            "Произошла ошибка при создании отчета.",
            reply_markup=get_main_keyboard()
        )

    return ConversationHandler.END


async def show_history(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    print(f"\nЗапрос истории от пользователя {user.id}")

    history_file = "history_requests.json"
    create_history_file()  # Создаем файл, если его нет

    try:
        with open(history_file, 'r', encoding='utf-8') as f:
            history_data = json.load(f)

        if not history_data:
            print("История запросов пуста")
            await update.message.reply_text("История запросов пуста.")
            return

        # Сортируем по дате (новые сначала)
        sorted_history = sorted(
            history_data.items(),
            key=lambda x: datetime.strptime(x[1], "%Y-%m-%d %H:%M:%S"),
            reverse=True
        )

        # Формируем сообщение
        message = "История запросов:\n\n"
        for org, date in sorted_history:
            message += f"📅 {date}\n🏢 {org}\n\n"

        print(f"Отправка истории ({len(sorted_history)} записей)")
        await update.message.reply_text(message)
    except Exception as e:
        print(f"❌ Ошибка при чтении истории: {str(e)}")
        await update.message.reply_text(f"Ошибка при чтении истории: {str(e)}")


async def extract_report(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    report_name = update.message.text.strip()
    print(f"\nЗапрос отчета от пользователя {user.id}: {report_name}")

    if not report_name:
        print("❌ Пустое название отчета")
        await update.message.reply_text("Пожалуйста, введите название отчета.")
        return WAITING_REPORT_NAME

    # Добавляем расширение, если его нет
    if not report_name.lower().endswith('.docx'):
        report_name += '.docx'

    reports_dir = "Reports"
    report_path = os.path.join(reports_dir, report_name)
    print(f"Поиск отчета: {report_path}")

    if os.path.exists(report_path):
        print(f"Отчет найден, отправка пользователю")
        with open(report_path, 'rb') as f:
            await update.message.reply_document(
                document=f,
                caption=f"Отчет '{report_name}'",
                reply_markup=get_main_keyboard()
            )
    else:
        print(f"❌ Отчет не найден")
        await update.message.reply_text(
            f"Отчет '{report_name}' не найден.",
            reply_markup=get_main_keyboard()
        )

    return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    print(f"\nОтмена действия пользователем {user.id}")
    await update.message.reply_text(
        "Действие отменено.",
        reply_markup=get_main_keyboard()
    )
    return ConversationHandler.END


def main():
    print("=" * 50)
    print("Запуск телеграм-бота")
    print("=" * 50)

    # Создаем Application и передаем токен
    try:
        application = Application.builder().token(TOKEN).build()
        print("✔ Бот успешно инициализирован")
    except Exception as e:
        print(f"❌ Ошибка при инициализации бота: {str(e)}")
        return

    # Добавляем обработчики
    conv_handler = ConversationHandler(
        entry_points=[MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message)],
        states={
            WAITING_FILE: [MessageHandler(filters.Document.FileExtension("docx"), handle_file)],
            WAITING_REPORT_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, extract_report)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )

    application.add_handler(CommandHandler('start', start))
    application.add_handler(conv_handler)

    # Запускаем бота
    print("\nБот запущен и ожидает сообщений...")
    print("Нажмите Ctrl+C для остановки")
    print("=" * 50)

    try:
        application.run_polling()
    except Exception as e:
        print(f"❌ Ошибка при работе бота: {str(e)}")


if __name__ == '__main__':
    main()