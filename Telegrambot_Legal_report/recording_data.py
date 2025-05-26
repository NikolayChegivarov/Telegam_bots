import os
from docx import Document
from extraction import extract_org_data
import json
from datetime import datetime


def create_folder():
    """Создает папку Reports в корневой директории, если её нет"""
    reports_dir = "Reports"
    if not os.path.exists(reports_dir):
        try:
            os.makedirs(reports_dir)
            print(f"✔ Создана папка: {reports_dir}")
        except Exception as e:
            print(f"❌ Ошибка при создании папки {reports_dir}: {str(e)}")
            return None
    return reports_dir


def create_history_file():
    """Создает файл history_requests.json в корневой директории, если его нет"""
    history_file = "history_requests.json"
    if not os.path.exists(history_file):
        try:
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump({}, f)  # Создаем пустой словарь
            print(f"✔ Создан файл истории запросов: {history_file}")
        except Exception as e:
            print(f"❌ Ошибка при создании файла истории: {str(e)}")
    return history_file


def write_to_history(org_name):
    """Записывает новую запись в файл истории"""
    history_file = "history_requests.json"
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    try:
        # Читаем существующие данные
        with open(history_file, 'r', encoding='utf-8') as f:
            history_data = json.load(f)

        # Добавляем новую запись
        history_data[org_name] = current_date

        # Записываем обновленные данные
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history_data, f, ensure_ascii=False, indent=4)

        print(f"✔ Добавлена запись в историю: {org_name} - {current_date}")
    except Exception as e:
        print(f"❌ Ошибка при записи в историю: {str(e)}")


def process_template(template_path, output_path, org_data):
    """Прямая запись данных в конкретные ячейки таблицы"""
    print(f"\n{'=' * 50}\nНачало обработки шаблона\n{'=' * 50}")

    # Проверка файла шаблона
    if not os.path.exists(template_path):
        print(f"❌ ОШИБКА: Файл шаблона не найден: {template_path}")
        return False

    try:
        doc = Document(template_path)
        print("✔ Шаблон успешно загружен")
    except Exception as e:
        print(f"❌ ОШИБКА при открытии шаблона: {str(e)}")
        return False

    # Проверка данных
    print("\nПроверка данных:")
    for key, value in org_data.items():
        print(f"{key}: {'✔' if value else '❌'} {value or 'Не найдено'}")

    # Обработка таблицы
    if not doc.tables:
        print("❌ В документе нет таблиц!")
        return False

    table = doc.tables[0]  # Первая таблица
    print(f"\nОбработка таблицы: {len(table.rows)} строк, {len(table.columns)} столбцов")

    try:
        # Ячейка Наименование (строка 1, столбец 2)
        name_cell = table.cell(0, 1)  # Индексы с 0
        org_name = org_data.get('Организация', 'Наименование не найдено')
        name_cell.text = org_name
        print(f"✔ Записано наименование в ячейку 1.2: {name_cell.text}")

        # Ячейка ОГРН (строка 2, столбец 2)
        ogrn_cell = table.cell(1, 1)  # Индексы с 0
        ogrn_cell.text = org_data.get('ОГРН', 'ОГРН не найден')
        print(f"✔ Записан ОГРН в ячейку 2.2: {ogrn_cell.text}")

        # Ячейка ИНН/КПП (строка 3, столбец 2)
        inn_kpp_cell = table.cell(2, 1)
        inn = org_data.get('ИНН', 'ИНН не найден')
        kpp = org_data.get('КПП', 'КПП не найден')
        inn_kpp_cell.text = f"{inn}/{kpp}"
        print(f"✔ Записан ИНН/КПП в ячейку 3.2: {inn_kpp_cell.text}")

    except IndexError:
        print("❌ ОШИБКА: Неправильная структура таблицы!")
        return False

    # Создаем файл истории (если его нет)
    create_history_file()

    # Записываем в историю (если название организации найдено)
    if org_name != 'Наименование не найдено':
        write_to_history(org_name)

    # Сохранение
    try:
        # Создаем папку Reports
        reports_dir = create_folder()
        if not reports_dir:
            return False

        # Формируем имя файла из названия организации
        clean_name = "".join(c for c in org_name if c.isalnum() or c in (' ', '_', '-'))
        clean_name = clean_name.strip()[:50]  # Ограничиваем длину имени файла
        if not clean_name:
            clean_name = "Отчет_без_названия"

        output_filename = f"{clean_name}.docx"
        full_output_path = os.path.join(reports_dir, output_filename)

        # Сохраняем файл
        doc.save(full_output_path)
        print(f"\n✔ Файл успешно сохранен: {full_output_path}")
        print(f"Размер файла: {os.path.getsize(full_output_path)} байт")
        return True
    except Exception as e:
        print(f"❌ ОШИБКА при сохранении: {str(e)}")
        return False


if __name__ == "__main__":
    source_file = 'контур_фокус.docx'
    template_file = 'шаблон.docx'
    output_file = 'НОВЫЙ_ФАЙЛ.docx'  # Это значение теперь не используется

    print("=" * 50)
    print("Начало обработки документов")
    print("=" * 50)

    # Извлекаем данные
    org_data = extract_org_data(source_file)

    # Обрабатываем шаблон (output_file теперь не используется)
    success = process_template(template_file, output_file, org_data)

    if success:
        print("\n✅ Обработка завершена успешно!")
    else:
        print("\n❌ В процессе обработки возникли ошибки!")

# import os
# from docx import Document
# from extraction import extract_org_data
#
#
# def create_folder():
#     """Создает папку Reports в корневой директории, если её нет"""
#     reports_dir = "Reports"
#     if not os.path.exists(reports_dir):
#         try:
#             os.makedirs(reports_dir)
#             print(f"✔ Создана папка: {reports_dir}")
#         except Exception as e:
#             print(f"❌ Ошибка при создании папки {reports_dir}: {str(e)}")
#             return None
#     return reports_dir
#
#
# def process_template(template_path, output_path, org_data):
#     """Прямая запись данных в конкретные ячейки таблицы"""
#     print(f"\n{'=' * 50}\nНачало обработки шаблона\n{'=' * 50}")
#
#     # Проверка файла шаблона
#     if not os.path.exists(template_path):
#         print(f"❌ ОШИБКА: Файл шаблона не найден: {template_path}")
#         return False
#
#     try:
#         doc = Document(template_path)
#         print("✔ Шаблон успешно загружен")
#     except Exception as e:
#         print(f"❌ ОШИБКА при открытии шаблона: {str(e)}")
#         return False
#
#     # Проверка данных
#     print("\nПроверка данных:")
#     for key, value in org_data.items():
#         print(f"{key}: {'✔' if value else '❌'} {value or 'Не найдено'}")
#
#     # Обработка таблицы
#     if not doc.tables:
#         print("❌ В документе нет таблиц!")
#         return False
#
#     table = doc.tables[0]  # Первая таблица
#     print(f"\nОбработка таблицы: {len(table.rows)} строк, {len(table.columns)} столбцов")
#
#     try:
#         # Ячейка Наименование (строка 1, столбец 2)
#         name_cell = table.cell(0, 1)  # Индексы с 0
#         org_name = org_data.get('Организация', 'Наименование не найдено')
#         name_cell.text = org_name
#         print(f"✔ Записано наименование в ячейку 1.2: {name_cell.text}")
#
#         # Ячейка ОГРН (строка 2, столбец 2)
#         ogrn_cell = table.cell(1, 1)  # Индексы с 0
#         ogrn_cell.text = org_data.get('ОГРН', 'ОГРН не найден')
#         print(f"✔ Записан ОГРН в ячейку 2.2: {ogrn_cell.text}")
#
#         # Ячейка ИНН/КПП (строка 3, столбец 2)
#         inn_kpp_cell = table.cell(2, 1)
#         inn = org_data.get('ИНН', 'ИНН не найден')
#         kpp = org_data.get('КПП', 'КПП не найден')
#         inn_kpp_cell.text = f"{inn}/{kpp}"
#         print(f"✔ Записан ИНН/КПП в ячейку 3.2: {inn_kpp_cell.text}")
#
#     except IndexError:
#         print("❌ ОШИБКА: Неправильная структура таблицы!")
#         return False
#
#     # Сохранение
#     try:
#         # Создаем папку Reports
#         reports_dir = create_folder()
#         if not reports_dir:
#             return False
#
#         # Формируем имя файла из названия организации
#         clean_name = "".join(c for c in org_name if c.isalnum() or c in (' ', '_', '-'))
#         clean_name = clean_name.strip()[:50]  # Ограничиваем длину имени файла
#         if not clean_name:
#             clean_name = "Отчет_без_названия"
#
#         output_filename = f"{clean_name}.docx"
#         full_output_path = os.path.join(reports_dir, output_filename)
#
#         # Сохраняем файл
#         doc.save(full_output_path)
#         print(f"\n✔ Файл успешно сохранен: {full_output_path}")
#         print(f"Размер файла: {os.path.getsize(full_output_path)} байт")
#         return True
#     except Exception as e:
#         print(f"❌ ОШИБКА при сохранении: {str(e)}")
#         return False
#
#
# if __name__ == "__main__":
#     source_file = 'контур_фокус.docx'
#     template_file = 'шаблон.docx'
#     output_file = 'НОВЫЙ_ФАЙЛ.docx'  # Это значение теперь не используется
#
#     print("=" * 50)
#     print("Начало обработки документов")
#     print("=" * 50)
#
#     # Извлекаем данные
#     org_data = extract_org_data(source_file)
#
#     # Обрабатываем шаблон (output_file теперь не используется)
#     success = process_template(template_file, output_file, org_data)
#
#     if success:
#         print("\n✅ Обработка завершена успешно!")
#     else:
#         print("\n❌ В процессе обработки возникли ошибки!")
