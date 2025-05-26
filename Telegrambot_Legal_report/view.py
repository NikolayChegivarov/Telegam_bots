from docx import Document
import os
from extraction import extract_org_data


def print_doc_structure(doc):
    """Выводит структуру документа в терминал"""
    print("\nСтруктура документа:")
    for table_idx, table in enumerate(doc.tables, 1):
        print(f"\nТаблица #{table_idx} - {len(table.rows)} строк, {len(table.columns)} столбцов")
        for row_idx, row in enumerate(table.rows, 1):
            for cell_idx, cell in enumerate(row.cells, 1):
                cell_text = cell.text.replace('\n', '\\n')
                print(f"Ячейка {row_idx}.{cell_idx}: '{cell_text}'")


def process_template(template_path, output_path, org_data):
    """Точечная вставка данных в конкретные ячейки"""
    print(f"\nОбработка шаблона: {template_path}")

    # Проверка файла
    if not os.path.exists(template_path):
        print(f"❌ Файл не найден: {template_path}")
        return False

    try:
        doc = Document(template_path)
        print("✔ Документ успешно загружен")
    except Exception as e:
        print(f"❌ Ошибка при открытии: {e}")
        return False

    # Выводим структуру документа перед изменениями
    print_doc_structure(doc)

    # Точечная вставка данных
    changes_made = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Вставляем ОГРН во вторую ячейку второй строки
                if cell.text.strip() == 'ОГРН:':
                    target_cell = row.cells[1]  # Соседняя ячейка справа
                    target_cell.text = org_data['ОГРН'] or 'ОГРН не найден'
                    print(f"\nВставка ОГРН в ячейку: '{target_cell.text}'")
                    changes_made = True

                # Вставляем ИНН/КПП во вторую ячейку третьей строки
                elif cell.text.strip() == 'ИНН/КПП:':
                    target_cell = row.cells[1]  # Соседняя ячейка справа
                    replacement = f"{org_data['ИНН'] or 'ИНН не найден'}/{org_data['КПП'] or 'КПП не найден'}"
                    target_cell.text = replacement
                    print(f"\nВставка ИНН/КПП в ячейку: '{target_cell.text}'")
                    changes_made = True

    if not changes_made:
        print("❌ Не найдены ячейки для вставки данных")
        return False

    # Сохранение файла
    try:
        doc.save(output_path)
        print(f"\n✔ Файл успешно сохранен: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Ошибка при сохранении: {e}")
        return False


if __name__ == "__main__":
    source_file = 'контур_фокус.docx'
    template_file = 'шаблон.docx'
    output_file = 'НОВЫЙ_ФАЙЛ.docx'

    print("=" * 50)
    print("Начало обработки документов")
    print("=" * 50)

    # Извлекаем данные
    org_data = extract_org_data(source_file)

    print("\nИзвлеченные данные:")
    for k, v in org_data.items():
        print(f"{k}: {v if v else 'Не найден'}")

    # Обрабатываем шаблон
    success = process_template(template_file, output_file, org_data)

    if success:
        print("\n✅ Обработка завершена успешно!")
    else:
        print("\n❌ В процессе обработки возникли ошибки!")